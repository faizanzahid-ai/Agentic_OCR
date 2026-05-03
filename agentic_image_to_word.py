"""
AGENTIC IMAGE-TO-WORD CONVERTER
--------------------------------
This script transforms the Phase 1 static tool into a Purely Agentic System,
addressing requirements 19 through 32 of Phase 2.

Required packages for full functionality:
pip install pytesseract python-docx watchdog pillow
"""

import os
import time
import json
import logging
from datetime import datetime


# Logging provides an audit trail of all actions the agent takes.
logging.basicConfig(
    filename="agent_safety.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Optional dependencies handling
try:
    from PIL import Image
    import pytesseract
    import platform
    # Auto-detect Tesseract binary on Windows in common install locations
    _found_tess = False
    if platform.system() == 'Windows':
        _tess_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Users\zahid\AppData\Local\Tesseract-OCR\tesseract.exe',
        ]
        for _p in _tess_paths:
            if os.path.exists(_p):
                pytesseract.pytesseract.tesseract_cmd = _p
                _found_tess = True
                break
    else:
        import shutil
        _found_tess = shutil.which('tesseract') is not None

    # Verify Tesseract actually works before enabling real OCR
    if _found_tess:
        try:
            pytesseract.get_tesseract_version()
            HAS_OCR = True
            print(f"[OK] Tesseract OCR engine found and ready.")
        except Exception:
            HAS_OCR = False
            print("[!] Tesseract binary found but failed to start. Falling back to simulation.")
    else:
        HAS_OCR = False
        print("[!] Tesseract OCR not installed. Install from: https://github.com/UB-Mannheim/tesseract/wiki")
        print("[!] Running in SIMULATION mode — text will be simulated, not extracted from image.")
except ImportError:
    HAS_OCR = False

try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    HAS_WATCHDOG = True
except ImportError:
    HAS_WATCHDOG = False


# 27. Memory & Context: 
class AgentMemory:
    def __init__(self, memory_file="agent_long_term_memory.json"):
        self.memory_file = memory_file
        self.short_term = {} # Context of current active session
        self.long_term = self._load_memory() # Historical user preferences

    def _load_memory(self):
        if os.path.exists(self.memory_file):
            with open(self.memory_file, 'r') as f:
                return json.load(f)
        return {"user_corrections": {}, "trusted_formats": []}

    def save_memory(self):
        with open(self.memory_file, 'w') as f:
            json.dump(self.long_term, f, indent=4)

    def learn_correction(self, original, corrected):
        # 20. Agentic System Concept: Learning (Updating knowledge base)
        self.long_term["user_corrections"][original] = corrected
        self.save_memory()
        logging.info(f"Learned new correction: '{original}' -> '{corrected}'")


#  Intelligence Layer: ML / Rules / LLMs
# -------------------------------------------------------------------------
class IntelligenceLayer:
    def __init__(self, memory: AgentMemory):
        self.memory = memory

    def interpret_image(self, image_path):
        """ 25. Operational Workflow: Interpret (Perception) """
        # Simulating a Vision ML model or OCR Engine
        logging.info(f"Interpreting image: {image_path}")
        print(f"[INTELLIGENCE] Extracting text and features from {image_path}...")
        
        if HAS_OCR:
            try:
                start_time = time.time()
                print(f"[AGENT] Scanning image... (this may take 10-20 seconds)")
                
                # Use optimized config for faster extraction (OEM 1 = Neural nets LSTM, PSM 3 = Auto page segmentation)
                custom_config = r'--oem 1 --psm 3'
                text = pytesseract.image_to_string(Image.open(image_path), config=custom_config)
                
                elapsed = time.time() - start_time
                if not text.strip():
                    print(f"[!] OCR finished in {elapsed:.1f}s but returned empty text.")
                    text = "[Agent Note: OCR produced no readable text. The image may be too blurry or dark.]"
                else:
                    print(f"[OCR] Success! Extracted {len(text.split())} words in {elapsed:.1f} seconds.")
            except Exception as e:
                logging.error(f"OCR Runtime Error: {e}")
                print(f"[ERROR] OCR failed at runtime: {e}")
                text = f"[Agent Error: OCR could not process this image. Reason: {e}]"
        else:
            # Graceful simulation fallback when Tesseract is not installed
            print("[SIMULATION] Tesseract not available — generating sample structured text.")
            text = (
                "QUARTERLY FINANCIAL REPORT\n"
                "Revenue increased by 15% this quarter.\n"
                "- Operating costs reduced by 8%\n"
                "- Net profit margin improved to 22%\n"
                "Key Highlights\n"
                "Performance exceeded all benchmarks set in Q1."
            )
        
        # Agent actively applies Long-term memory corrections to raw data
        for wrong, right in self.memory.long_term["user_corrections"].items():
            text = text.replace(wrong, right)
            
        return text

    def determine_formatting(self, text):
        """ 20. Agentic System Concept: Decision-making (Multi-rule intelligence) """
        # Agent uses multiple weighted rules instead of static hardcoded logic.
        # This overcomes the core limitation from Phase 1 (Req 19).
        lines = text.split('\n')
        formatted_blocks = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # --- RULE ENGINE (26. Intelligence Layer) ---
            length     = len(stripped)
            word_count = len(stripped.split())
            digit_ratio = sum(c.isdigit() for c in stripped) / max(length, 1)
            alpha_ratio = sum(c.isalpha() for c in stripped) / max(length, 1)

            # Rule 1: Short all-caps or title-case → strong heading signal
            is_heading = (
                (length < 60 and stripped.isupper()) or
                (length < 50 and stripped.istitle() and word_count <= 8)
            )

            # Rule 2: Bullet / list item
            is_list = stripped.startswith(('-', '*', '•', '·'))

            # Rule 3: Very short fragments with high digit ratio → noisy OCR
            is_noise = length < 4 or digit_ratio > 0.6

            # --- CONFIDENCE SCORING (32. Explainability) ---
            if is_noise:
                confidence = 0.45          # Flag for human review
            elif is_heading:
                confidence = 0.95          # High confidence classification
            elif is_list:
                confidence = 0.90          # List items are well-defined
            elif alpha_ratio > 0.7 and word_count >= 5:
                confidence = 0.92          # Clean prose paragraph
            elif word_count >= 3:
                confidence = 0.85          # Reasonable text line
            else:
                confidence = 0.75          # Short / ambiguous line

            # Determine block type label for explainability
            if is_noise:
                block_type = "noise"
            elif is_heading:
                block_type = "heading"
            elif is_list:
                block_type = "list_item"
            else:
                block_type = "paragraph"

            formatted_blocks.append({
                "text":       stripped,
                "is_heading": is_heading,
                "is_list":    is_list,
                "block_type": block_type,
                "confidence": round(confidence, 2),
            })
        return formatted_blocks


# -------------------------------------------------------------------------
#  Agent Architecture: Flow: Input -> Processing -> Decision -> Action -> Feedback
#  Agent Type Selection: Learning agent
# -------------------------------------------------------------------------
class IntelliDocAgent:
    def __init__(self):
        self.memory = AgentMemory()
        self.intelligence = IntelligenceLayer(self.memory)
        # 30. Ethical Agent Design: Privacy (Ensuring local processing)
        logging.info("IntelliDoc Agent initialized. Data stays on local machine.")

    def run_agentic_workflow(self, file_path):
        """ 25. Operational Workflow: Observe -> Interpret -> Decide -> Act -> Learn """
        print(f"\n{'='*50}\n[AGENT START] Processing: {os.path.basename(file_path)}")
        
        # Step 2: Interpret
        raw_text = self.intelligence.interpret_image(file_path)
        
        # Step 3: Decide
        print("[AGENT] Deciding on layout and semantic structure...")
        formatted_blocks = self.intelligence.determine_formatting(raw_text)
        
        # 31. Risk Assessment: Preventing Over-automation and Incorrect Decisions
        low_confidence_blocks = [b for b in formatted_blocks if b['confidence'] < 0.8]
        
        # Step 4: Act
        print("[AGENT] Taking Action: Constructing formatted Word Document...")
        docx_path = self.generate_word_document(file_path, formatted_blocks)
        
        # Step 5: Learn & Update (Feedback Loop)
        corrections = self.human_in_the_loop_review(docx_path, low_confidence_blocks)
        
        # Apply corrections to the CURRENT session and re-generate if needed
        if corrections:
            print("[AGENT] Applying your exact corrections to the current document...")
            for block in formatted_blocks:
                for wrong, right in corrections.items():
                    block["text"] = block["text"].replace(wrong, right)
            # Re-save the final version
            self.generate_word_document(file_path, formatted_blocks, is_final=True)

    def generate_word_document(self, file_path, formatted_blocks, is_final=False):
        """ Action Component. """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        suffix = "_Final" if is_final else "_Draft"
        output_path = file_path.rsplit('.', 1)[0] + f"_Agentic_{timestamp}{suffix}.docx"
        
        if HAS_DOCX:
            doc = Document()
            for block in formatted_blocks:
                if block["block_type"] == "noise":
                    continue  # 31. Risk: Skip unreadable OCR noise
                p = doc.add_paragraph()
                run = p.add_run(block["text"])
                if block["is_heading"]:
                    run.bold = True
                    run.font.size = Pt(16)
                elif block["is_list"]:
                    run.font.size = Pt(11)
                    p.style = doc.styles["List Bullet"] if "List Bullet" in [s.name for s in doc.styles] else p.style
                # Append confidence footnote comment to low-confidence lines
                if block["confidence"] < 0.8:
                    doc.add_paragraph(f"[⚠ Low confidence ({block['confidence']}) — Please verify: '{block['text'][:60]}']")
            doc.save(output_path)
            logging.info(f"Successfully generated: {output_path}")
        else:
            print("[!] python-docx not installed. Simulating document generation.")

        print(f"[ACTION SUCCESS] File saved to -> {output_path}")
        return output_path

    # -------------------------------------------------------------------------
    #  Autonomy Level: Semi-autonomy (User validates final output)
    #  Human-in-the-Loop: Human controls the system to prevent misuse.
    # -------------------------------------------------------------------------
    def human_in_the_loop_review(self, docx_path, low_confidence_blocks):
        print("\n" + "="*50)
        print("  [HUMAN-IN-THE-LOOP] Feedback & Review Phase")
        print("="*50)

        if low_confidence_blocks:
            print(f"\n[!] Agent flagged {len(low_confidence_blocks)} text block(s) with low confidence:")
            for i, b in enumerate(low_confidence_blocks, 1):
                snippet = b['text'][:60] + ('...' if len(b['text']) > 60 else '')
                print(f"  [{i}] type={b['block_type']:10s} | conf={b['confidence']} | '{snippet}'")
        else:
            print("\n[OK] All blocks classified with high confidence.")

        corrections_made = 0

        # --- Question 1 ---
        print("\n" + "-"*40)
        print("Q1. Are you satisfied with the overall document output?")
        q1 = input("    Enter (y)es / (n)o: ").strip().lower()
        if q1 == 'n':
            print("    [NOTE] Noted. Your feedback is logged for agent improvement.")
            logging.info("User dissatisfied with output.")
        else:
            print("    [OK] Great! Satisfaction recorded.")
            logging.info("User satisfied with output.")

        # --- Question 2 ---
        print("\n" + "-"*40)
        print("Q2. Did the agent misread any word? (OCR correction)")
        print("    Example: type '1,l' OR just type the wrong word.")
        q2_input = input("    Enter correction or press Enter to skip: ").strip()
        
        q2_wrong, q2_right = None, None
        if ',' in q2_input:
            q2_wrong, q2_right = [x.strip() for x in q2_input.split(',', 1)]
        elif q2_input:
            q2_wrong = q2_input
            q2_right = input(f"    [AGENT] You typed '{q2_wrong}'. What should it be corrected to? ").strip()
        
        if q2_wrong and q2_right:
            self.memory.learn_correction(q2_wrong, q2_right)
            print(f"    [LEARN] Saved: '{q2_wrong}' -> '{q2_right}' in Long-Term Memory.")
            corrections_made += 1
        else:
            print("    [SKIP] No OCR correction needed.")

        # --- Question 3 ---
        print("\n" + "-"*40)
        print("Q3. Was any heading incorrectly formatted as body text (or vice versa)?")
        q3 = input("    Enter (y)es / (n)o: ").strip().lower()
        if q3 == 'y':
            print("    Please enter the exact text that was misclassified:")
            bad_line = input("    Text: ").strip()
            print("    How should it be classified? (heading / paragraph / list)")
            correct_type = input("    Correct type: ").strip().lower()
            logging.info(f"User correction: '{bad_line}' should be '{correct_type}'")
            self.memory.long_term.setdefault("layout_corrections", {})[bad_line] = correct_type
            self.memory.save_memory()
            print(f"    [LEARN] Stored layout correction for future runs.")
            corrections_made += 1
        else:
            print("    [OK] Formatting was correct.")

        # --- Question 4 ---
        print("\n" + "-"*40)
        print("Q4. Do you want to teach the agent another word correction?")
        q4_input = input("    Enter correction or press Enter to skip: ").strip()
        
        q4_wrong, q4_right = None, None
        if ',' in q4_input:
            q4_wrong, q4_right = [x.strip() for x in q4_input.split(',', 1)]
        elif q4_input:
            q4_wrong = q4_input
            q4_right = input(f"    [AGENT] You typed '{q4_wrong}'. What should it be corrected to? ").strip()
            
        if q4_wrong and q4_right:
            self.memory.learn_correction(q4_wrong, q4_right)
            print(f"    [LEARN] Saved: '{q4_wrong}' -> '{q4_right}' in Long-Term Memory.")
            corrections_made += 1
        else:
            print("    [SKIP] No additional corrections.")

        # --- Question 5: Final Approval ---
        print("\n" + "-"*40)
        print("Q5. Final step — do you approve this document for use?")
        q5 = input("    Enter (y)es to approve, (n)o to reject: ").strip().lower()
        if q5 == 'y':
            logging.info(f"Document APPROVED by user: {docx_path}")
            print(f"\n  [APPROVED] Document '{os.path.basename(docx_path)}' has been approved.")
        else:
            logging.warning(f"Document REJECTED by user: {docx_path}")
            print(f"\n  [REJECTED] Document marked as rejected. Please re-run with a clearer image.")

        # Summary
        print("\n" + "="*50)
        print(f"  Feedback Summary:")
        print(f"  - Corrections taught to agent : {corrections_made}")
        print(f"  - Document status             : {'APPROVED' if q5 == 'y' else 'REJECTED'}")
        print(f"  - Long-term memory saved to   : agent_long_term_memory.json")
        print("="*50)

        # Return corrections to apply them to the current document
        applied_corrections = {}
        if q2_wrong and q2_right:
            applied_corrections[q2_wrong] = q2_right
        if q4_wrong and q4_right:
            applied_corrections[q4_wrong] = q4_right
            
        return applied_corrections


# -------------------------------------------------------------------------
# Gap Analysis: Addressing lack of autonomy.
# Agentic Vision: Transform Tool -> Agent, Reactive -> Proactive.
# -------------------------------------------------------------------------
if HAS_WATCHDOG:
    class ProactiveFolderMonitor(FileSystemEventHandler):
        def __init__(self, agent):
            self.agent = agent

        def on_created(self, event):
            # Proactive monitoring: System acts autonomously when environment changes
            if not event.is_directory and event.src_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                print(f"\n[OBSERVE] Agent perceived new file in environment: {event.src_path}")
                self.agent.run_agentic_workflow(event.src_path)

def start_proactive_agent():
    watch_folder = "./Agent_Inbox"
    if not os.path.exists(watch_folder):
        os.makedirs(watch_folder)

    agent = IntelliDocAgent()

    if HAS_WATCHDOG:
        observer = Observer()
        observer.schedule(ProactiveFolderMonitor(agent), watch_folder, recursive=False)
        print("\n" + "*"*60)
        print("* AGENTIC VISION REALIZED (Req 22): Proactive Mode Active  *")
        print(f"* Monitoring Folder: {os.path.abspath(watch_folder)}")
        print("* Drop a .jpg or .png image here to trigger the Agent.       *")
        print("*"*60)
        observer.start()
        try:
            while True:
                time.sleep(1)
        except KeyboardInterrupt:
            observer.stop()
        observer.join()
    else:
        print("[!] 'watchdog' library not found. Falling back to Reactive Execution.")
        print("[!] For full Agentic Vision (Req 22), run: pip install watchdog")
        # Reactive simulation
        sample_img = os.path.join(watch_folder, "test_document.jpg")
        with open(sample_img, 'w') as f: f.write("dummy") 
        print(f"\n[OBSERVE] Manually passing {sample_img} to Agent...")
        agent.run_agentic_workflow(sample_img)

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        agent = IntelliDocAgent()
        agent.run_agentic_workflow(sys.argv[1])
    else:
        start_proactive_agent()
